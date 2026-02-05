# app_full.py
import os
import traceback
from dotenv import load_dotenv
from typing import List, Optional
from datetime import datetime, timedelta
import uuid
import sqlite3
import secrets
from functools import wraps

# --- Environment Variable Loading ---
load_dotenv()
os.environ['GROQ_API_KEY'] = os.getenv("GROQ_API_KEY")
os.environ['GOOGLE_API_KEY'] = os.getenv("GOOGLE_API_KEY")
# Optional Tesseract path
TESSERACT_CMD = os.getenv("TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe")

# --- FastAPI Imports ---
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Depends, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel

# --- Security helpers ---
from werkzeug.security import generate_password_hash, check_password_hash

# --- Core Logic Imports ---
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from PIL import Image
import pytesseract
import speech_recognition as sr
from pydub import AudioSegment


from translation_utils import TranslationManager
# Configure Tesseract
pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# --- LangChain Imports ---
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document as LCDocument
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain.memory import ConversationBufferWindowMemory
from langchain.chains import ConversationalRetrievalChain
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage

# --- Initialization ---
app = FastAPI(title="Full Integrated Legal Chat (FastAPI)")

# --- CORS Configuration ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # restrict in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- DATABASE SETUP (SQLite) ---
DB_PATH = "conversations.db"

def init_database():
    """Initialize SQLite database for conversation storage"""
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    
    # Create users table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id TEXT PRIMARY KEY,
            username TEXT UNIQUE NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create conversations table (with user_id)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS conversations (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            title TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            message_count INTEGER DEFAULT 0,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Create messages table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id TEXT NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (conversation_id) REFERENCES conversations (id)
        )
    ''')
    
    # Create sessions table for token management
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sessions (
            token TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            expires_at TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    conn.commit()
    conn.close()

# Init DB immediately
init_database()

# --- Helper DB functions ---
def db_connect():
    return sqlite3.connect(DB_PATH, check_same_thread=False)

# --- Authentication dependency for FastAPI ---
def get_current_user(authorization: Optional[str] = Header(None)):
    """Dependency to fetch user_id from Authorization header 'Bearer <token>'"""
    if not authorization:
        raise HTTPException(status_code=401, detail="No authentication token provided")
    token = authorization
    if token.startswith("Bearer "):
        token = token[7:]
    conn = db_connect()
    cur = conn.cursor()
    cur.execute('''
        SELECT user_id FROM sessions WHERE token = ? AND (expires_at IS NULL OR expires_at > CURRENT_TIMESTAMP)
    ''', (token,))
    row = cur.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    return row[0]

# --- Pydantic Models ---
class RegisterModel(BaseModel):
    username: str
    email: str
    password: str

class LoginModel(BaseModel):
    email: str
    password: str

class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    question: str
    # --- CHANGED: Added conversation_id to the request model ---
    conversation_id: Optional[str] = None
    chat_history: List[ChatMessage] = []
    session_id: str = "default"

class ResetRequest(BaseModel):
    session_id: str = "default"

# --- Global Variables & Model Loading ---
translation_manager = TranslationManager()
llm_text, llm_vision, BASE_RETRIEVER, embeddings = None, None, None, None
try:
    print("Initializing models and vector store...")
    groq_api_key = os.getenv("GROQ_API_KEY")
    if groq_api_key:
        llm_text = ChatGroq(groq_api_key=groq_api_key, model_name="llama-3.3-70b-versatile")
    else:
        print("Warning: GROQ_API_KEY not set - llm_text will be None")

    # Google vision LLM (optional)
    google_key = os.getenv("GOOGLE_API_KEY")
    if google_key:
        llm_vision = ChatGoogleGenerativeAI(model="gemini-pro-vision", temperature=0.3)
    else:
        print("Warning: GOOGLE_API_KEY not set - llm_vision will be None")

    print("Initializing local embedding model...")
    model_name = "sentence-transformers/all-MiniLM-L6-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}
    embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )
    print("Local embedding model loaded.")
    # Load FAISS vectorstore (if present)
    try:
        db = FAISS.load_local("my_vector_store", embeddings, allow_dangerous_deserialization=True)
        BASE_RETRIEVER = db.as_retriever(search_type="similarity", search_kwargs={"k": 4})
        print("Vector store loaded and retriever available.")
    except Exception as e:
        print("No vector store found or failed to load. You can upload docs to create session retrievers.")
        BASE_RETRIEVER = None

except Exception:
    print("FATAL ERROR during initialization:")
    traceback.print_exc()
    llm_text, llm_vision, BASE_RETRIEVER, embeddings = None, None, None, None

SESSION_RETRIEVERS = {}
DEFAULT_SESSION_ID = "default"
if BASE_RETRIEVER:
    SESSION_RETRIEVERS[DEFAULT_SESSION_ID] = BASE_RETRIEVER

# --- Utility functions for conversations & messages ---
def generate_conversation_title(first_message):
    """Generate a conversation title - attempt to use llm_text, fallback to first words"""
    try:
        if llm_text:
            prompt = f"""Generate a short, descriptive title (maximum 6 words) for a conversation that starts with this question: "{first_message[:200]}"
Return only the title, nothing else."""
            response = llm_text.invoke(prompt)
            title = response.content.strip().replace('"', '').replace("'", "")
            return title[:50]
    except Exception:
        pass
    words = first_message.split()[:4]
    return ' '.join(words) + ('...' if len(first_message.split()) > 4 else '')

def save_conversation_message(conversation_id, role, content):
    conn = db_connect()
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO messages (conversation_id, role, content) VALUES (?, ?, ?)
    ''', (conversation_id, role, content))
    cursor.execute('''
        UPDATE conversations SET updated_at = CURRENT_TIMESTAMP,
            message_count = (SELECT COUNT(*) FROM messages WHERE conversation_id = ?)
        WHERE id = ?
    ''', (conversation_id, conversation_id))
    conn.commit()
    conn.close()

def create_new_conversation(first_message, user_id):
    conversation_id = str(uuid.uuid4())
    title = generate_conversation_title(first_message)
    conn = db_connect()
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO conversations (id, user_id, title) VALUES (?, ?, ?)
    ''', (conversation_id, user_id, title))
    conn.commit()
    conn.close()
    return conversation_id

def get_conversation_history(conversation_id):
    conn = db_connect()
    cursor = conn.cursor()
    cursor.execute('''
        SELECT role, content FROM messages WHERE conversation_id = ? ORDER BY timestamp ASC
    ''', (conversation_id,))
    messages = [{'role': r[0], 'content': r[1]} for r in cursor.fetchall()]
    conn.close()
    return messages

def get_all_conversations(user_id):
    conn = db_connect()
    cursor = conn.cursor()
    cursor.execute('''
        SELECT id, title, created_at, updated_at, message_count FROM conversations
        WHERE user_id = ? ORDER BY updated_at DESC
    ''', (user_id,))
    conversations = []
    for row in cursor.fetchall():
        conversations.append({
            'id': row[0],
            'title': row[1],
            'created_at': row[2],
            'updated_at': row[3],
            'message_count': row[4]
        })
    conn.close()
    return conversations

def delete_conversation(conversation_id, user_id):
    conn = db_connect()
    cursor = conn.cursor()
    cursor.execute('SELECT user_id FROM conversations WHERE id = ?', (conversation_id,))
    row = cursor.fetchone()
    if not row or row[0] != user_id:
        conn.close()
        return False
    cursor.execute('DELETE FROM messages WHERE conversation_id = ?', (conversation_id,))
    cursor.execute('DELETE FROM conversations WHERE id = ?', (conversation_id,))
    conn.commit()
    conn.close()
    return True

def verify_conversation_ownership(conversation_id, user_id):
    conn = db_connect()
    cursor = conn.cursor()
    cursor.execute('SELECT user_id FROM conversations WHERE id = ?', (conversation_id,))
    row = cursor.fetchone()
    conn.close()
    return bool(row and row[0] == user_id)

# --- Related questions helper ---
def get_related_questions(user_question, bot_answer):
    related_prompt = f"""
You are an smart legal assistant.
Based on the user's last question and your answer, suggest 4 short, clear related legal questions
that the user might want to ask next.
Only output the questions as a numbered list, without extra text.

User Question: {user_question}
Your Answer: {bot_answer}
"""
    try:
        if not llm_text:
            return []
        response = llm_text.invoke(related_prompt)
        suggestions = response.content.strip().split("\n")
        return [q.strip("0123456789. ").strip() for q in suggestions if q.strip()]
    except Exception:
        return []


# --- LANGUAGE ENDPOINT ---
@app.get("/languages")
def get_supported_languages():
    """Returns a list of supported languages for the UI."""
    if not translation_manager.is_available():
        raise HTTPException(status_code=503, detail="Translation service is unavailable")
    return translation_manager.get_supported_languages_for_ui()


# --- AUTH ENDPOINTS ---
@app.post("/register")
def register(body: RegisterModel):
    username = body.username
    email = body.email
    password = body.password
    if not all([username, email, password]):
        raise HTTPException(status_code=400, detail="Username, email, and password are required")
    try:
        user_id = str(uuid.uuid4())
        password_hash = generate_password_hash(password)
        conn = db_connect()
        cur = conn.cursor()
        cur.execute('''
            INSERT INTO users (id, username, email, password_hash) VALUES (?, ?, ?, ?)
        ''', (user_id, username, email, password_hash))
        conn.commit()
        conn.close()
        return {"message": "User registered successfully", "user_id": user_id}
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="Username or email already exists")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/login")
def login(body: LoginModel):
    email = body.email
    password = body.password
    if not all([email, password]):
        raise HTTPException(status_code=400, detail="Email and password are required")
    try:
        conn = db_connect()
        cur = conn.cursor()
        cur.execute('SELECT id, password_hash, username FROM users WHERE email = ?', (email,))
        row = cur.fetchone()
        if not row or not check_password_hash(row[1], password):
            conn.close()
            raise HTTPException(status_code=401, detail="Invalid email or password")
        token = secrets.token_urlsafe(32)
        user_id = row[0]
        username = row[2]
        # optional expiry: e.g., 7 days
        expires_at = (datetime.utcnow() + timedelta(days=7)).isoformat()
        cur.execute('INSERT INTO sessions (token, user_id, expires_at) VALUES (?, ?, ?)', (token, user_id, expires_at))
        conn.commit()
        conn.close()
        return {"token": token, "user_id": user_id, "username": username}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/logout")
def logout(authorization: Optional[str] = Header(None)):
    if not authorization:
        raise HTTPException(status_code=401, detail="No authentication token provided")
    token = authorization
    if token.startswith("Bearer "):
        token = token[7:]
    try:
        conn = db_connect()
        cur = conn.cursor()
        cur.execute('DELETE FROM sessions WHERE token = ?', (token,))
        conn.commit()
        conn.close()
        return {"message": "Logged out successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/verify")
def verify(current_user_id: str = Depends(get_current_user)):
    try:
        conn = db_connect()
        cur = conn.cursor()
        cur.execute('SELECT username, email FROM users WHERE id = ?', (current_user_id,))
        row = cur.fetchone()
        conn.close()
        if row:
            return {"valid": True, "user_id": current_user_id, "username": row[0], "email": row[1]}
        else:
            raise HTTPException(status_code=401, detail="User not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- Conversation CRUD Endpoints ---
@app.get("/conversations")
def get_conversations(current_user_id: str = Depends(get_current_user)):
    try:
        conversations = get_all_conversations(current_user_id)
        return {"conversations": conversations}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/conversations/{conversation_id}")
def get_conversation(conversation_id: str, current_user_id: str = Depends(get_current_user)):
    if not verify_conversation_ownership(conversation_id, current_user_id):
        raise HTTPException(status_code=403, detail="Unauthorized access to conversation")
    messages = get_conversation_history(conversation_id)
    return {"messages": messages}

@app.post("/conversations/new")
def create_conversation_endpoint(current_user_id: str = Depends(get_current_user)):
    conversation_id = str(uuid.uuid4())
    conn = db_connect()
    cur = conn.cursor()
    cur.execute('INSERT INTO conversations (id, user_id, title) VALUES (?, ?, ?)', (conversation_id, current_user_id, "New Chat"))
    conn.commit()
    conn.close()
    return {"conversation_id": conversation_id}

@app.delete("/conversations/{conversation_id}")
def delete_conversation_endpoint(conversation_id: str, current_user_id: str = Depends(get_current_user)):
    try:
        if delete_conversation(conversation_id, current_user_id):
            return {"message": "Conversation deleted successfully"}
        else:
            raise HTTPException(status_code=403, detail="Conversation not found or unauthorized")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- Chat Endpoint ---
@app.post("/chat")
def chat(request: ChatRequest, current_user_id: str = Depends(get_current_user)):
    if not llm_text:
        raise HTTPException(status_code=500, detail="LLM not initialized")
    
    question = request.question
    session_id = request.session_id or DEFAULT_SESSION_ID
    conversation_id = request.conversation_id # Get ID from request body

    # --- CHANGED: Correct logic to handle existing and new conversations ---
    if not conversation_id:
        # If no ID is provided, create a new conversation
        conversation_id = create_new_conversation(question, current_user_id)
    else:
        # If an ID is provided, verify the user owns it
        if not verify_conversation_ownership(conversation_id, current_user_id):
            raise HTTPException(status_code=403, detail="Unauthorized access to conversation")

    # Save user message to the correct conversation
    save_conversation_message(conversation_id, 'user', question)

    # Build memory from chat_history field
    memory = ConversationBufferWindowMemory(k=3, memory_key="chat_history", return_messages=True)
    # --- ADDED: Load previous messages from DB for better context ---
    chat_history = get_conversation_history(conversation_id)
    for message in chat_history:
        if message['role'] == 'user':
            memory.chat_memory.add_user_message(message['content'])
        else:
            memory.chat_memory.add_ai_message(message['content'])

    retriever = SESSION_RETRIEVERS.get(session_id, BASE_RETRIEVER)
    prompt_template = """
<s>[INST] You are an expert legal chatbot. Your primary goal is to provide clear, accurate, and well-structured information.
**Instructions:**
- Analyze the user's question and the provided context carefully.
- Structure your answer using Markdown for readability.
- Use numbered or bulleted lists for enumerations (like legal clauses or steps).
- Use **bold text** to highlight key legal terms, parties (like "landlord" and "tenant"), and important concepts.
- Keep your tone professional and direct.
CONTEXT: {context}
CHAT HISTORY: {chat_history}
QUESTION: {question}
ANSWER (in Markdown format):
</s>[INST]
"""
    prompt = PromptTemplate(template=prompt_template, input_variables=['context', 'question', 'chat_history'])
    qa_chain = ConversationalRetrievalChain.from_llm(
        llm=llm_text,
        memory=memory,
        retriever=retriever,
        combine_docs_chain_kwargs={'prompt': prompt}
    )
    try:
        result = qa_chain.invoke(input=question)
        answer = result.get("answer", "Sorry, I could not find an answer.")
        related_questions = get_related_questions(question, answer)
        save_conversation_message(conversation_id, 'assistant', answer)
        return {"answer": answer, "related_questions": related_questions, "conversation_id": conversation_id}
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

# --- Vision / OCR Chat Endpoint ---
@app.post("/chat_vision")
async def chat_vision(
    question: str = Form(...), 
    image: UploadFile = File(...), 
    current_user_id: str = Depends(get_current_user), 
    conversation_id: Optional[str] = Form(None),
    language: str = Form("en") # Add language field
):
    if not image or not question:
        raise HTTPException(status_code=400, detail="Image and question are required")
    
    target_lang = language or "en"
    
    # --- Conversation Management ---
    if not conversation_id:
        conversation_id = create_new_conversation(question, current_user_id)
    else:
        if not verify_conversation_ownership(conversation_id, current_user_id):
            raise HTTPException(status_code=403, detail="Unauthorized access to conversation")
    
    # Save user's original message
    save_conversation_message(conversation_id, 'user', question)

    try:
        # --- Language Translation (Input) ---
        if target_lang != 'en':
            translated_question = translation_manager.translate_text(question, target_lang, 'en')
            if not translated_question:
                raise HTTPException(status_code=500, detail="Failed to translate question")
        else:
            translated_question = question
            
        image_bytes = await image.read()
        pil_image = Image.open(BytesIO(image_bytes)).convert("RGB")
        
        # OCR extraction
        extracted_text = pytesseract.image_to_string(pil_image)
        if not extracted_text.strip():
            answer = "I could not find any readable text in the image. Please try a clearer image."
            save_conversation_message(conversation_id, 'assistant', answer)
            return {"answer": answer, "related_questions": [], "conversation_id": conversation_id}
            
        # Build prompt for text LLM using extracted text
        ocr_prompt_template = """
You are a highly intelligent legal analysis assistant. Your task is to answer the user's question based *only* on the context provided below, which has been extracted from an image using OCR.

**Context (Text extracted from image):**
---
{context}
---

**User's Question:**
{question}

**Instructions:**
- Analyze the extracted text to answer the question accurately.
- If the text is insufficient to answer the question, state that the information is not present in the provided text.
- Structure your answer clearly using Markdown.

**Answer:
"""
        prompt = PromptTemplate(template=ocr_prompt_template, input_variables=['context', 'question'])
        
        if not llm_text:
            raise HTTPException(status_code=500, detail="Text LLM not available")

        # Get answer in English
        chain = prompt | llm_text
        response = chain.invoke({"context": extracted_text, "question": translated_question})
        answer_en = response.content
        related_questions_en = get_related_questions(translated_question, answer_en)

        # --- Language Translation (Output) ---
        if target_lang != 'en':
            final_answer = translation_manager.translate_text(answer_en, 'en', target_lang) or answer_en
            final_related_questions = [
                translation_manager.translate_text(q, 'en', target_lang) or q for q in related_questions_en
            ]
        else:
            final_answer = answer_en
            final_related_questions = related_questions_en
        
        # Save translated assistant message
        save_conversation_message(conversation_id, 'assistant', final_answer)
        return {
            "answer": final_answer, 
            "related_questions": final_related_questions, 
            "conversation_id": conversation_id
        }
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to process image query: {str(e)}")

from typing import List # Make sure this is imported at the top

# --- Document Upload Endpoint (index into vector store session) ---
@app.post("/upload_document")
async def upload_document(
    session_id: str = Form(DEFAULT_SESSION_ID), 
    files: List[UploadFile] = File(...), # UPDATE: Changed to accept a list of files
    current_user_id: str = Depends(get_current_user)
):
    if not embeddings:
        raise HTTPException(status_code=500, detail="Embeddings model not initialized")
    if not files:
        raise HTTPException(status_code=400, detail="No files part")

    try:
        all_texts = []
        processed_filenames = []

        # UPDATE: Loop through the list of files
        for file in files:
            raw_text = ""
            file_ext = file.filename.split('.')[-1].lower()
            file_content = await file.read()
            
            if file_ext == "pdf":
                reader = PdfReader(BytesIO(file_content))
                for page in reader.pages:
                    raw_text += page.extract_text() or ""
            elif file_ext == "docx":
                doc = Document(BytesIO(file_content))
                for para in doc.paragraphs:
                    raw_text += (para.text or "") + "\n"
            elif file_ext == "txt":
                raw_text = file_content.decode('utf-8')
            else:
                # You can choose to skip unsupported files or raise an error
                print(f"Skipping unsupported file type: {file.filename}")
                continue # Skip to the next file

            if raw_text:
                all_texts.append(raw_text)
                processed_filenames.append(file.filename)

        if all_texts:
            # Combine text from all processed files
            combined_text = "\n\n".join(all_texts)
            text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            texts = text_splitter.split_text(combined_text)
            
            documents = [LCDocument(page_content=t) for t in texts]
            
            # Create or update the vector store with documents from all files
            temp_vector_store = FAISS.from_documents(documents, embeddings)
            SESSION_RETRIEVERS[session_id] = temp_vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 4})
            
            return {"message": f"Successfully processed {len(processed_filenames)} files: {', '.join(processed_filenames)}."}
        else:
            raise HTTPException(status_code=400, detail="No text could be extracted from the uploaded files.")
            
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to process files: {str(e)}")
# --- Reset session retriever ---
@app.post("/reset")
def reset(request: ResetRequest, current_user_id: str = Depends(get_current_user)):
    SESSION_RETRIEVERS[request.session_id] = BASE_RETRIEVER
    return {"message": "Conversation reset successfully."}

# --- Audio Transcription Endpoint ---
@app.post("/transcribe")
async def transcribe_audio(audio: UploadFile = File(...), current_user_id: str = Depends(get_current_user)):
    if not audio:
        raise HTTPException(status_code=400, detail="No audio file part")
    recognizer = sr.Recognizer()
    try:
        audio_content = await audio.read()
        # browser might send webm/ogg - convert to wav using pydub
        sound = AudioSegment.from_file(BytesIO(audio_content))
        wav_io = BytesIO()
        sound.export(wav_io, format="wav")
        wav_io.seek(0)
        with sr.AudioFile(wav_io) as source:
            audio_data = recognizer.record(source)
        text = recognizer.recognize_google(audio_data)
        return {"transcription": text}
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Could not transcribe audio: {str(e)}")

# --- Run using uvicorn externally (uvicorn app_full:app --reload --port 5001) ---
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app_full:app", host="0.0.0.0", port=5001, reload=True)
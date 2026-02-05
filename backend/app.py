import os
import traceback
from dotenv import load_dotenv
import base64
from io import BytesIO
import json
import uuid
from datetime import datetime
import sqlite3
import secrets
from functools import wraps

# --- OCR Library Import ---
import pytesseract

# --- Explicitly load environment variables ---
load_dotenv()
os.environ['GROQ_API_KEY'] = os.getenv("GROQ_API_KEY")

from flask import Flask, request, jsonify
from flask_cors import CORS
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
from werkzeug.security import generate_password_hash, check_password_hash

from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document as LCDocument
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain.memory import ConversationBufferWindowMemory
from langchain.chains import ConversationalRetrievalChain

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.messages import HumanMessage

# Audio Processing Imports
import speech_recognition as sr
from pydub import AudioSegment

# Add this line to specify the Tesseract path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# --- Initialization ---
app = Flask(__name__)
CORS(app)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', secrets.token_hex(32))

# --- Database Setup for Conversation Storage ---
def init_database():
    """Initialize SQLite database for conversation storage"""
    conn = sqlite3.connect('conversations.db')
    cursor = conn.cursor()
    
    # Create users table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id TEXT PRIMARY KEY,
            username TEXT UNIQUE NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create conversations table (with user_id)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS conversations (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            title TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            message_count INTEGER DEFAULT 0,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Create messages table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id TEXT NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (conversation_id) REFERENCES conversations (id)
        )
    ''')
    
    # Create sessions table for token management
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sessions (
            token TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            expires_at TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    conn.commit()
    conn.close()

# Initialize database on startup
init_database()

# --- Authentication Middleware ---
def require_auth(f):
    """Decorator to require authentication for endpoints"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        token = request.headers.get('Authorization')
        
        if not token:
            return jsonify({"error": "No authentication token provided"}), 401
        
        # Remove 'Bearer ' prefix if present
        if token.startswith('Bearer '):
            token = token[7:]
        
        user_id = verify_token(token)
        if not user_id:
            return jsonify({"error": "Invalid or expired token"}), 401
        
        # Add user_id to request context
        request.user_id = user_id
        return f(*args, **kwargs)
    
    return decorated_function

def verify_token(token):
    """Verify authentication token and return user_id"""
    conn = sqlite3.connect('conversations.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT user_id FROM sessions 
        WHERE token = ? AND (expires_at IS NULL OR expires_at > CURRENT_TIMESTAMP)
    ''', (token,))
    
    result = cursor.fetchone()
    conn.close()
    
    return result[0] if result else None

# --- Authentication Endpoints ---
@app.route('/register', methods=['POST'])
def register():
    """Register a new user"""
    data = request.json
    username = data.get('username')
    email = data.get('email')
    password = data.get('password')
    
    if not all([username, email, password]):
        return jsonify({"error": "Username, email, and password are required"}), 400
    
    try:
        user_id = str(uuid.uuid4())
        password_hash = generate_password_hash(password)
        
        conn = sqlite3.connect('conversations.db')
        cursor = conn.cursor()
        
        cursor.execute('''
            INSERT INTO users (id, username, email, password_hash)
            VALUES (?, ?, ?, ?)
        ''', (user_id, username, email, password_hash))
        
        conn.commit()
        conn.close()
        
        return jsonify({"message": "User registered successfully", "user_id": user_id})
    
    except sqlite3.IntegrityError:
        return jsonify({"error": "Username or email already exists"}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/login', methods=['POST'])
def login():
    """Login user and return authentication token"""
    data = request.json
    email = data.get('email')
    password = data.get('password')
    
    if not all([email, password]):
        return jsonify({"error": "Email and password are required"}), 400
    
    try:
        conn = sqlite3.connect('conversations.db')
        cursor = conn.cursor()
        
        cursor.execute('SELECT id, password_hash, username FROM users WHERE email = ?', (email,))
        user = cursor.fetchone()
        
        if not user or not check_password_hash(user[1], password):
            conn.close()
            return jsonify({"error": "Invalid email or password"}), 401
        
        # Create session token
        token = secrets.token_urlsafe(32)
        user_id = user[0]
        username = user[2]
        
        cursor.execute('''
            INSERT INTO sessions (token, user_id)
            VALUES (?, ?)
        ''', (token, user_id))
        
        conn.commit()
        conn.close()
        
        return jsonify({
            "token": token,
            "user_id": user_id,
            "username": username
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/logout', methods=['POST'])
@require_auth
def logout():
    """Logout user by invalidating token"""
    token = request.headers.get('Authorization')
    if token.startswith('Bearer '):
        token = token[7:]
    
    try:
        conn = sqlite3.connect('conversations.db')
        cursor = conn.cursor()
        
        cursor.execute('DELETE FROM sessions WHERE token = ?', (token,))
        
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Logged out successfully"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/verify', methods=['GET'])
@require_auth
def verify():
    """Verify if token is still valid"""
    try:
        conn = sqlite3.connect('conversations.db')
        cursor = conn.cursor()
        
        cursor.execute('SELECT username, email FROM users WHERE id = ?', (request.user_id,))
        user = cursor.fetchone()
        conn.close()
        
        if user:
            return jsonify({
                "valid": True,
                "user_id": request.user_id,
                "username": user[0],
                "email": user[1]
            })
        else:
            return jsonify({"valid": False}), 401
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- Conversation Management Functions ---
def generate_conversation_title(first_message):
    """Generate a title for conversation based on first message"""
    try:
        if llm_text:
            prompt = f"""Generate a short, descriptive title (maximum 6 words) for a conversation that starts with this question: "{first_message[:200]}"
            
            Return only the title, nothing else."""
            
            response = llm_text.invoke(prompt)
            title = response.content.strip().replace('"', '').replace("'", "")
            return title[:50]  # Limit to 50 characters
    except:
        pass
    
    # Fallback: use first few words of the message
    words = first_message.split()[:4]
    return ' '.join(words) + ('...' if len(first_message.split()) > 4 else '')

def save_conversation_message(conversation_id, role, content):
    """Save a message to the database"""
    conn = sqlite3.connect('conversations.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO messages (conversation_id, role, content)
        VALUES (?, ?, ?)
    ''', (conversation_id, role, content))
    
    # Update conversation updated_at and message_count
    cursor.execute('''
        UPDATE conversations 
        SET updated_at = CURRENT_TIMESTAMP,
            message_count = (SELECT COUNT(*) FROM messages WHERE conversation_id = ?)
        WHERE id = ?
    ''', (conversation_id, conversation_id))
    
    conn.commit()
    conn.close()

def create_new_conversation(first_message, user_id):
    """Create a new conversation and return its ID"""
    conversation_id = str(uuid.uuid4())
    title = generate_conversation_title(first_message)
    
    conn = sqlite3.connect('conversations.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO conversations (id, user_id, title)
        VALUES (?, ?, ?)
    ''', (conversation_id, user_id, title))
    
    conn.commit()
    conn.close()
    
    return conversation_id

def get_conversation_history(conversation_id):
    """Get all messages for a specific conversation"""
    conn = sqlite3.connect('conversations.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT role, content FROM messages 
        WHERE conversation_id = ? 
        ORDER BY timestamp ASC
    ''', (conversation_id,))
    
    messages = [{'role': row[0], 'content': row[1]} for row in cursor.fetchall()]
    conn.close()
    
    return messages

def get_all_conversations(user_id):
    """Get list of all conversations for a specific user"""
    conn = sqlite3.connect('conversations.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT id, title, created_at, updated_at, message_count 
        FROM conversations 
        WHERE user_id = ?
        ORDER BY updated_at DESC
    ''', (user_id,))
    
    conversations = []
    for row in cursor.fetchall():
        conversations.append({
            'id': row[0],
            'title': row[1],
            'created_at': row[2],
            'updated_at': row[3],
            'message_count': row[4]
        })
    
    conn.close()
    return conversations

def delete_conversation(conversation_id, user_id):
    """Delete a conversation and all its messages (only if owned by user)"""
    conn = sqlite3.connect('conversations.db')
    cursor = conn.cursor()
    
    # Verify ownership
    cursor.execute('SELECT user_id FROM conversations WHERE id = ?', (conversation_id,))
    result = cursor.fetchone()
    
    if not result or result[0] != user_id:
        conn.close()
        return False
    
    cursor.execute('DELETE FROM messages WHERE conversation_id = ?', (conversation_id,))
    cursor.execute('DELETE FROM conversations WHERE id = ?', (conversation_id,))
    
    conn.commit()
    conn.close()
    return True

def verify_conversation_ownership(conversation_id, user_id):
    """Verify if a conversation belongs to the user"""
    conn = sqlite3.connect('conversations.db')
    cursor = conn.cursor()
    
    cursor.execute('SELECT user_id FROM conversations WHERE id = ?', (conversation_id,))
    result = cursor.fetchone()
    conn.close()
    
    return result and result[0] == user_id

# --- Global Variables & Model Loading ---
llm_text, BASE_RETRIEVER, embeddings = None, None, None
try:
    groq_api_key = os.getenv("GROQ_API_KEY")
    llm_text = ChatGroq(groq_api_key=groq_api_key, model_name="llama-3.3-70b-versatile")

    print("Initializing local embedding model...")
    model_name = "sentence-transformers/all-MiniLM-L6-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}
    embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )
    print("âœ… Local embedding model loaded.")
    
    db = FAISS.load_local("my_vector_store", embeddings, allow_dangerous_deserialization=True)
    BASE_RETRIEVER = db.as_retriever(search_type="similarity", search_kwargs={"k": 4})
    print("âœ… Models and vector store initialized successfully.")
except Exception as e:
    print(f"âŒ FATAL ERROR during initialization:")
    traceback.print_exc()
    llm_text, BASE_RETRIEVER, embeddings = None, None, None

SESSION_RETRIEVERS = {}
DEFAULT_SESSION_ID = "default"
if BASE_RETRIEVER:
    SESSION_RETRIEVERS[DEFAULT_SESSION_ID] = BASE_RETRIEVER

# --- Helper Functions ---
def get_related_questions(user_question, bot_answer):
    related_prompt = f"""
    You are an Indian Penal Code (IPC) legal assistant.
    Based on the user's last question and your answer, suggest 4 short, clear related legal questions
    that the user might want to ask next.
    Only output the questions as a numbered list, without extra text.

    User Question: {user_question}
    Your Answer: {bot_answer}
    """
    try:
        response = llm_text.invoke(related_prompt)
        suggestions = response.content.strip().split("\n")
        return [q.strip("0123456789. ").strip() for q in suggestions if q.strip()]
    except Exception:
        return []

# --- Conversation API Endpoints ---
@app.route('/conversations', methods=['GET'])
@require_auth
def get_conversations():
    """Get all conversations list for authenticated user"""
    try:
        conversations = get_all_conversations(request.user_id)
        return jsonify({"conversations": conversations})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/conversations/<conversation_id>', methods=['GET'])
@require_auth
def get_conversation(conversation_id):
    """Get a specific conversation with all messages"""
    try:
        # Verify ownership
        if not verify_conversation_ownership(conversation_id, request.user_id):
            return jsonify({"error": "Unauthorized access to conversation"}), 403
        
        messages = get_conversation_history(conversation_id)
        return jsonify({"messages": messages})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/conversations/<conversation_id>', methods=['DELETE'])
@require_auth
def delete_conversation_endpoint(conversation_id):
    """Delete a specific conversation"""
    try:
        if delete_conversation(conversation_id, request.user_id):
            return jsonify({"message": "Conversation deleted successfully"})
        else:
            return jsonify({"error": "Conversation not found or unauthorized"}), 403
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/conversations/new', methods=['POST'])
@require_auth
def create_conversation():
    """Create a new conversation"""
    try:
        conversation_id = str(uuid.uuid4())
        conn = sqlite3.connect('conversations.db')
        cursor = conn.cursor()
        
        cursor.execute('''
            INSERT INTO conversations (id, user_id, title)
            VALUES (?, ?, ?)
        ''', (conversation_id, request.user_id, "New Chat"))
        
        conn.commit()
        conn.close()
        
        return jsonify({"conversation_id": conversation_id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- Chat Endpoint ---
@app.route('/chat', methods=['POST'])
@require_auth
def chat():
    if not llm_text:
        return jsonify({"error": "LLM not initialized"}), 500
    
    data = request.json
    question = data.get('question')
    conversation_id = data.get('conversation_id')
    session_id = data.get('session_id', DEFAULT_SESSION_ID)
    
    if not question:
        return jsonify({"error": "Question is required"}), 400
    
    # If no conversation_id provided, create a new one for this user
    if not conversation_id:
        conversation_id = create_new_conversation(question, request.user_id)
    else:
        # Verify ownership
        if not verify_conversation_ownership(conversation_id, request.user_id):
            return jsonify({"error": "Unauthorized access to conversation"}), 403
    
    # Get conversation history from database
    chat_history = get_conversation_history(conversation_id)
    
    # Save user message
    save_conversation_message(conversation_id, 'user', question)
    
    retriever = SESSION_RETRIEVERS.get(session_id, BASE_RETRIEVER)
    memory = ConversationBufferWindowMemory(
        k=3, memory_key="chat_history", return_messages=True
    )
    
    # Load chat history into memory
    for message in chat_history:
        if message['role'] == 'user':
            memory.chat_memory.add_user_message(message['content'])
        else:
            memory.chat_memory.add_ai_message(message['content'])
    
    prompt_template = """
    <s>[INST] You are an expert legal chatbot. Your primary goal is to provide clear, accurate, and well-structured information.
    **Instructions:**
    - Analyze the user's question and the provided context carefully.
    - Structure your answer using Markdown for readability.
    - Use numbered or bulleted lists for enumerations (like legal clauses or steps).
    - Use **bold text** to highlight key legal terms, parties (like "landlord" and "tenant"), and important concepts.
    - Keep your tone professional and direct.
    CONTEXT: {context}
    CHAT HISTORY: {chat_history}
    QUESTION: {question}
    ANSWER (in Markdown format):
    </s>[INST]
    """
    
    prompt = PromptTemplate(
        template=prompt_template,
        input_variables=['context', 'question', 'chat_history']
    )
    
    qa_chain = ConversationalRetrievalChain.from_llm(
        llm=llm_text,
        memory=memory,
        retriever=retriever,
        combine_docs_chain_kwargs={'prompt': prompt}
    )
    
    try:
        result = qa_chain.invoke(input=question)
        answer = result.get("answer", "Sorry, I could not find an answer.")
        related_questions = get_related_questions(question, answer)
        
        # Save assistant message
        save_conversation_message(conversation_id, 'assistant', answer)
        
        return jsonify({
            "answer": answer,
            "related_questions": related_questions,
            "conversation_id": conversation_id
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- Vision Chat Endpoint ---
@app.route('/chat_vision', methods=['POST'])
@require_auth
def chat_vision():
    """Handles image queries by first extracting text with OCR, then analyzing with LLM"""
    if 'image' not in request.files or 'question' not in request.form:
        return jsonify({"error": "Image and question are required"}), 400

    question = request.form['question']
    conversation_id = request.form.get('conversation_id')
    image_file = request.files['image']

    # If no conversation_id provided, create a new one for this user
    if not conversation_id:
        conversation_id = create_new_conversation(question, request.user_id)
    else:
        # Verify ownership
        if not verify_conversation_ownership(conversation_id, request.user_id):
            return jsonify({"error": "Unauthorized access to conversation"}), 403

    try:
        # Save user message
        save_conversation_message(conversation_id, 'user', question)
        
        # Step 1: Extract text from image using OCR
        image = Image.open(image_file.stream)
        extracted_text = pytesseract.image_to_string(image)

        if not extracted_text.strip():
            answer = "I could not find any readable text in the image. Please try a clearer image."
            save_conversation_message(conversation_id, 'assistant', answer)
            return jsonify({
                "answer": answer,
                "related_questions": [],
                "conversation_id": conversation_id
            })
            
        # Step 2: Build prompt for text LLM using extracted text
        ocr_prompt_template = """
        You are a highly intelligent legal analysis assistant. Your task is to answer the user's question based *only* on the context provided below, which has been extracted from an image using OCR.

        **Context (Text extracted from image):**
        ---
        {context}
        ---

        **User's Question:**
        {question}

        **Instructions:**
        - Analyze the extracted text to answer the question accurately.
        - If the text is insufficient to answer the question, state that the information is not present in the provided text.
        - Structure your answer clearly using Markdown.
        
        **Answer:**
        """
        
        prompt = PromptTemplate(
            template=ocr_prompt_template,
            input_variables=['context', 'question']
        )
        
        # Step 3: Create chain and invoke text LLM
        chain = prompt | llm_text
        response = chain.invoke({
            "context": extracted_text,
            "question": question
        })
        
        # Step 4: Generate related questions and save response
        answer = response.content
        related_questions = get_related_questions(question, answer)
        
        # Save assistant message
        save_conversation_message(conversation_id, 'assistant', answer)

        return jsonify({
            "answer": answer, 
            "related_questions": related_questions,
            "conversation_id": conversation_id
        })
    
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Failed to process image query: {str(e)}"}), 500

# --- Document Upload Endpoint ---
@app.route('/upload_document', methods=['POST'])
@require_auth
def upload_file():
    if not embeddings:
        return jsonify({"error": "Embeddings model not initialized"}), 500
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    session_id = request.form.get('session_id', f"{request.user_id}_{DEFAULT_SESSION_ID}")
    try:
        raw_text = ""
        file_ext = file.filename.split('.')[-1].lower()
        if file_ext == "pdf":
            reader = PdfReader(file)
            for page in reader.pages:
                raw_text += page.extract_text() or ""
        elif file_ext == "docx":
            doc = Document(BytesIO(file.read()))
            for para in doc.paragraphs:
                raw_text += (para.text or "") + "\n"
        elif file_ext == "txt":
            raw_text = file.read().decode('utf-8')
        if raw_text:
            text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            texts = text_splitter.split_text(raw_text)
            documents = [LCDocument(page_content=t) for t in texts]
            temp_vector_store = FAISS.from_documents(documents, embeddings)
            SESSION_RETRIEVERS[session_id] = temp_vector_store.as_retriever(
                search_type="similarity", search_kwargs={"k": 4}
            )
            return jsonify({"message": f"File '{file.filename}' processed successfully."})
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Failed to process file: {str(e)}"}), 500
    return jsonify({"error": "File type not supported"}), 400

@app.route('/reset', methods=['POST'])
@require_auth
def reset():
    session_id = request.json.get('session_id', f"{request.user_id}_{DEFAULT_SESSION_ID}")
    SESSION_RETRIEVERS[session_id] = BASE_RETRIEVER
    return jsonify({"message": "Conversation reset successfully."})

@app.route('/transcribe', methods=['POST'])
@require_auth
def transcribe_audio():
    if 'audio' not in request.files:
        return jsonify({"error": "No audio file part"}), 400
    audio_file = request.files['audio']
    recognizer = sr.Recognizer()
    try:
        sound = AudioSegment.from_file(audio_file, format="webm")
        wav_io = BytesIO()
        sound.export(wav_io, format="wav")
        wav_io.seek(0)
        with sr.AudioFile(wav_io) as source:
            audio_data = recognizer.record(source)
        text = recognizer.recognize_google(audio_data)
        return jsonify({"transcription": text})
    except Exception as e:
        return jsonify({"error": f"Could not transcribe audio: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5001, debug=True)